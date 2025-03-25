import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document
from langchain.embeddings import HuggingFaceBgeEmbeddings
from langchain.vectorstores import FAISS
from ChatGLM3_main.cli_demo_thirteen import main as chat
from rank_bm25 import BM25Okapi 
import jieba
doc_path = "D:/Age_copy/law_doc"

documnets = []
# def load_docx(file_path):
#     doc = Document(file_path)
#     text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
#     return text

# for file_name in os.listdir(doc_path):
#     if file_name.endswith(".docx"):
#         file_path = os.path.join(doc_path,file_name)
#         try:
#             text = load_docx(file_path)
#             documnets.append(text)
#         except Exception as e:
#             print(f"读取{file_name}失败:{e}")
# text_splitter = RecursiveCharacterTextSplitter(chunk_size=128,chunk_overlap=50)
# docs = text_splitter.create_documents(documnets)
# print(f"成功载入{len(docs)}个文档块")

# # Step 3 Convert Documents inot Embeddings
# embdding_model = HuggingFaceBgeEmbeddings(model_name="sensenova/piccolo-base-zh")
# vector_store = FAISS.from_documents(docs,embdding_model)
# vector_store.save_local("faiss_index")
# print("文档已存入FAISS向量数据库")

# # Step4 Build BM25 index
# bm25_corpus = [list(jieba.cut(doc.page_content)) for doc in docs]
# bm25 = BM25Okapi(bm25_corpus)

# # Step5 RAG combine BM25 with Embedding Similiarity Search
# def rag_ask(query):
#      max_lenght=90
#      # BM25 keyword search
#      tokenized_query = list(jieba.cut(query))
#      bm25_scores = bm25.get_scores(tokenized_query)

#      #  Get BM25 Top 5 documents
#      bm25_top_k = 3
#      bm25_top_docs = sorted(zip(bm25_scores,docs),key=lambda x:x[0],reverse=True)[:bm25_top_k]

#      # From BM25 outcome select top 3 docs for Embedding Search
#      retrieved_docs = [doc[1] for doc in bm25_top_docs]
#      embedding_retrieved_docs = vector_store.similarity_search(query,k=2)  # vetor search

#      # Combine BM25 and Embedding Search Results
#      combined_docs = {doc.page_content:doc for doc in retrieved_docs + embedding_retrieved_docs}
#      retrieved_docs = list(combined_docs.values())[:2]

#      context = "\n\n".join([doc.page_content for doc in retrieved_docs])
#      prompt = f"以下是相关资料：\n{context}\n\n用户问题:{query}\n\n请根据以下资料回答用户问题："
#      response = chat(prompt,1)
#      return response

import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document
from langchain.embeddings import HuggingFaceBgeEmbeddings
from langchain.vectorstores import FAISS
from ChatGLM3_main.cli_demo_thirteen import main as chat
from rank_bm25 import BM25Okapi 
import jieba
doc_path = "D:/Age_copy/law_doc"

documnets = []
def load_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return text

for file_name in os.listdir(doc_path):
    if file_name.endswith(".docx"):
        file_path = os.path.join(doc_path,file_name)
        try:
            text = load_docx(file_path)
            documnets.append(text)
        except Exception as e:
            print(f"读取{file_name}失败:{e}")
text_splitter = RecursiveCharacterTextSplitter(chunk_size=512,chunk_overlap=50)
docs = text_splitter.create_documents(documnets)
print(f"成功载入{len(docs)}个文档块")

# Step 3 Convert Documents inot Embeddings
embdding_model = HuggingFaceBgeEmbeddings(model_name="sensenova/piccolo-base-zh")
vector_store = FAISS.from_documents(docs,embdding_model)
vector_store.save_local("faiss_index")
print("文档已存入FAISS向量数据库")

# Step4 Build BM25 index
bm25_corpus = [list(jieba.cut(doc.page_content)) for doc in docs]
bm25 = BM25Okapi(bm25_corpus)

# Step5 RAG combine BM25 with Embedding Similiarity Search
def rag_ask(query):
     max_lenght=100
     # BM25 keyword search
     tokenized_query = list(jieba.cut(query))
     bm25_scores = bm25.get_scores(tokenized_query)

     #  Get BM25 Top 5 documents
     bm25_top_k = 5
     bm25_top_docs = sorted(zip(bm25_scores,docs),key=lambda x:x[0],reverse=True)[:bm25_top_k]

     # From BM25 outcome select top 3 docs for Embedding Search
     retrieved_docs = [doc[1] for doc in bm25_top_docs]
     embedding_retrieved_docs = vector_store.similarity_search(query,k=3)  # vetor search

     # Combine BM25 and Embedding Search Results
     combined_docs = {doc.page_content:doc for doc in retrieved_docs + embedding_retrieved_docs}
     retrieved_docs = list(combined_docs.values())[:3]

     context = "\n\n".join([doc.page_content for doc in retrieved_docs])
     prompt = f"以下是相关资料：\n{context}\n\n用户问题:{query}\n\n请根据以下资料回答用户问题："
     #response = chat(prompt,1)
     return prompt